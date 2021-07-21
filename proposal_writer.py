from docx import Document
from datetime import date
from docx.shared import Pt, Inches,RGBColor
import sys
from tkinter import *
from tkinter.filedialog import asksaveasfilename ,askopenfilename
import os
from openpyxl import Workbook
from openpyxl import load_workbook
import re



print('excel to py file loaded...')

#sys.path.append('C:\Users\Owner\Desktop\Estimating model 1.0.7.9')

#set global file name
file = ''
def set_file(_file):
    global file
    file = _file
    print(file)

estimate_log = {}

#Phrases used in proposal
rail_height = ['36"','42"','34"-38"','custom']
post_type = ['2-1/2 x 4" Base shoe','2-3/8" x 2-3/8" square aluminum posts', '2-3/8" x 2-3/8" Series 500 aluminum slotted Post','1.5" Sch 40','1.5" x 1.5" Square','1" x 2" Rectangular','custom' ]

mounting_detail = ["fascia mounted to front of deck framing using PRO's Fascia brackets", 'fascia mounted to front of deck framing using steel angle iron (angle iron installed by others)',
                    'mounted directly to deck framing using engineered lags','mounted to top of deck surface using rubber gasket and 5x5 baseplate','fascia mounted to welded knife plates (knife plates by others)',
                   'fascia mounted to angle aluminum brakets attached to halfens','mounted to stringer using 3" x 8" slotted angle base plate','1 line', '2 line','3 line','custom']

top_rail = ['Top rail profile 200', 'Top rail profile 375','Top rail profile 400','CL Laurence 1" x 1-5/16" SS  Top rail Profile','No top rail','core mounted grabrail','baseplate mounted grabrail','wall mounted grabrail','custom']

bottom_rail = ['bottom rail profile 200','with CR Laurence SS cladding','bottom rail profile 100','bottom rail profile 500','Without a Bottom rail','to be mounted directly to walls at stairs','to be mounted directly to aluminum posts at stairs',
               'to be mounted directly into core hole pockets','to be mounted directly to surface','custom']

infill = ['5/8" x 5/8" picket infill','1/4 laminated Tempered glass infill','3/8 laminated Tempered glass infill','1/2 laminated Tempered glass infill','9/16 laminated Tempered glass infill','1/8" SS Cable Infill','3/16" SS Cable infill','quikset grout',
          '4"x4" aluminum baseplate','decorative grabrail brackets','custom']

spacing = ['',"2'","3'","4'","5'","6'",'custom']

rail_type = ['Picket Guardrail','Glass Guardrail','Cable Guardrail','Grab rail','custom']


def search_estimates():
    '''search feature for estimate archive (est_log.txt)
        - prints results for full search term, and split words. 
    '''
    os.chdir('C:/Users/Owner/Desktop/Estimating model 1.0.7.9')
    tag = input("what are you searching for? : ")
    word_list = tag.split(' ')
    file = open('est_log.txt','r+')
    res_dict = {}
    for line in file:
        for word in word_list:
            if word.lower() in line.lower():
                res = re.findall(r"\d\d\d\d?\d",line)
                if word in res_dict:
                    res_dict[word].append(res[0])
                else:
                    res_dict[word] = []
                    res_dict[word].append(res[0])
    for key in res_dict:
        print('Search Tag:',"'",key,"'\n",'Estimate Numbers: ',res_dict[key],'\n')


def append_description(number,list_to_append,description_list,input_string):
    #generic function for return description to append phrase to list or ask for user input
    if description_list[number] == "custom":
        user_input = input(f"Need input for {input_string}")
        list_to_append.append(user_input)
    else:
        list_to_append.append(description_list[number])


def return_description(set_height,set_post_type,set_mount_type,set_toprail,set_bottomrail,set_infill,set_spacing,set_type):
    #takes index numbers and returns a list of phrases
    rail_description = []
    append_description(set_height,rail_description,rail_height,'Please set custom Railing height \n: ')
    append_description(set_post_type,rail_description,post_type,'Please set post type/ grabrail type \n: ')
    append_description(set_mount_type,rail_description,mounting_detail,'Please set mounting detail or # of grabrail lines \n: ')
    append_description(set_toprail,rail_description,top_rail,'Please set top rail or grabrail mounting (core/post/wall) \n: ')
    append_description(set_bottomrail,rail_description,bottom_rail,'Please set bottom rail or grabrail mounting (core/post/wall)\n:')
    append_description(set_infill,rail_description,infill,'Please set railing infill or grabrail mounting detail (grout/bp/bracket) \n: ')
    append_description(set_spacing,rail_description,spacing,'Please set post spacing \n: ')
    append_description(set_type,rail_description,rail_type,'Please set infill type (picket/glass/cable) \n: ')
    return rail_description



class Excel_to_py:
    # object used to interface with excel estimating model

    def __init__(self,file,workbook,note_sheet):
        self.file = file
        self.workbook = workbook
        note_sheet = workbook[note_sheet]
        self.note_sheet = note_sheet

    def return_job_info(self):
        #return customer info
        job_info = {}
        info_order = ['customer_name','company_name','contact_info','business_address','job_address','job_name']
        i = 0
        for row in self.note_sheet.iter_rows(min_row=1,min_col=2,max_col=2,values_only=True):
            res = ''.join(map(str,row))
            if res != 'None' and i < 6:
                job_info[info_order[i]] = res
                i += 1
        return job_info

    def return_lf(self):
        #return LF for each section
        section_lf= []
        for row in self.note_sheet.iter_rows(min_row=2,max_row=6,min_col=4,max_col=4,values_only=True):
            res = ''.join(map(str,row))
            if res != "None":
                if res != "0":
                    res = round(float(res),0)
                    section_lf.append(int(res))
                else:
                    section_lf.append('NA')          
        return section_lf

    def return_lfprice(self):
        # return price per LF for each section
        section_lfprice= []
        for row in self.note_sheet.iter_rows(min_row=2,max_row=6,min_col=5,max_col=5,values_only=True):
            res = ''.join(map(str,row))
            if res != 'None':
                if res != '0':
                    res = round(float(res),0)
                    section_lfprice.append(int(res))
                else:
                    section_lfprice.append('NA') 
        return section_lfprice

    def return_section_details(self,num=0):
        # return section detail numbers in the excel estimate
        sections = self.return_lf()
        section = []
        for row in self.note_sheet.iter_rows(min_row=2,min_col=(7+num),max_col=(7+num),values_only=True):
                res = ''.join(map(str,row))
                if res != 'None':
                    section.append(int(res))
        return section

    def return_area_name(self,num=0):
        # return section name
        area_names = ''
        for row in self.note_sheet.iter_rows(min_row=(2+num),max_row=(2+num),min_col=3,max_col=3,values_only=True):
            print(row)
            res = ''.join(map(str,row))
            if res != 'None':
                area_name = str(res)
                return area_name
            else:
                return 'None'
     
    def return_rep(self):
        # return Salesman
        for row in self.note_sheet.iter_rows(min_row=13,max_row=13,min_col=2,max_col=2,values_only=True):
            rep = ''.join(map(str,row))
            if rep.lower() == 'jag':
                return 'jag'
            elif rep.lower() == 'dave':
                return 'dave'
            else:
                return 'jag'
    

def write_proposal(instance):
    print('start making proposal...')
    #getting instance values for job lf and lf prices, rep
    bid_lf = instance.return_lf()
    bid_lfprice = instance.return_lfprice()
    job_info = instance.return_job_info()
    rep = instance.return_rep()
    estimate_total = 0

    #for saving to est_log
    estimate_details = []
    
    #load estimate number
    est_num_file = open('estimate_number.txt','r+')
    estimate_number = int(est_num_file.read())
    est_num_file.close()
    
    # Document Variables
    today = date.today()
    todays_date = today.strftime("%m/%d/%Y")
        
    #Document Setup
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    #margin setup
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(3.4)
    paragraph_format.space_after = Pt(0)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(.5)
        section.bottom_margin = Inches(.4)
        section.left_margin = Inches(.8)
        section.right_margin = Inches(.8)

    #header setup
    header = document.sections[0].header
    h1 = header.paragraphs[0].add_run("\tPrecision Rail of Oregon, LLC\n").font.size = Pt(36)
    h2 = header.paragraphs[0].add_run("10735 SE Foster RD").font.size = Pt(14)
    header.paragraphs[0].add_run('\t').add_picture('Alumarail_logo.png', width = Inches(1.8))
    h3 = header.paragraphs[0].add_run("\tPortland, Oregon 97266").font.size = Pt(14)
    h1.italic = True
    h2.italic = True
    h3.italic = True
    h1.bold = True
    h2.bold = True
    h3.bold = True
    header.paragraphs[0].style.font.color.rgb = RGBColor(0, 96, 43)

    #footer setup
    footer = document.sections[0].footer
    footer.paragraphs[0].add_run('Phone: 503-512-5353\tFax: 503-668-8968\twww.Precisionrail.com').font.size = Pt(14)
    footer.paragraphs[0].style.font.color.rgb = RGBColor(0, 96, 43)

    # start Document
    est =document.add_paragraph()
    est.add_run('\t'*10 + 'Estimate #: {}-C'.format(estimate_number)).bold=True
    
    #customer/job info
    for item in job_info.keys():
        estimate_details.append(job_info[item])
        document.add_paragraph(str(job_info[item]))
    document.add_paragraph('')
    
    document.add_paragraph('Dear ' + job_info['customer_name'] + ',\n')
    p1 = document.add_paragraph()
    p1.style = document.styles['Normal']
    
    p1.add_run('Precision Rail of Oregon is pleased to provide the following proposal for: ')
    p1.add_run(job_info['job_name'] + ', BUDGET Rev-0 \n\n').bold = True
    p1.add_run('Items furnished by Precision Rail of Oregon: Submittal drawings, engineering, materials, and installation.\n\n').bold = True
    p1.add_run('Submittals:').bold = True
    p1.add_run(' Pricing includes 1 submittal based off plans and 1 revision once corrections are received from GC. Any Additional revisions to be billed at 145.00 per hour plus materials and handling. \n\n')

    #section write up
    for num in range(len(bid_lf)):

        ''' for each area, if its lf != NA...
            - find area name, find section decriptions,append to estimate, '''
        
        if bid_lf[num] != 'NA':
            bid_area_total = bid_lf[num] * bid_lfprice[num]
            bid_area_name = instance.return_area_name(num)
            section_index_array = instance.return_section_details(num)
            finished_descriptions = return_description(section_index_array[0],section_index_array[1],section_index_array[2],section_index_array[3],
            section_index_array[4],section_index_array[5],section_index_array[6],section_index_array[7]) 
            estimate_details.append(finished_descriptions)
            p1.add_run(f"Bid Item - {finished_descriptions[0]} Tall {finished_descriptions[7]} ({bid_area_name})\n").bold = True
            p1.add_run(" {} {}. {}, {} with {}. Posts spacing to be evenly spaced and not exceed {} as per engineering and customer request. Support blocking by others. Standard color (Black, Bronze, White). ".format(finished_descriptions[1],finished_descriptions[2],
                                                                                                                                                                                                                         finished_descriptions[3],finished_descriptions[4],finished_descriptions[5],finished_descriptions[6]))    
            if finished_descriptions[7] == 'Grab rail':
                p1.add_run('Handrails are all ADA Compliant.')
            p1.add_run('\n\n')
            p1.add_run('\t'*6 + 'Sub Total {} LF @ ${}.00 per LF = ${}.00*\n\n'.format(str(bid_lf[num]),str(bid_lfprice[num]),str(bid_area_total))).bold = True
            estimate_total += bid_area_total

    #price total
    p1.add_run('\n'*3)
    p1.add_run('\t'*10 + '     Total = {}.00*\n\n\n'.format(str(estimate_total))).bold = True
    p1.add_run('\t\t*This price quote is valid for 30 days from the date of this document*\n\n').italic = True

    #typical Assumptions and exclusions for contract
    p1.add_run('Assumptions\n').bold = True
    p1.add_run('The following assumptions were made in support of this estimate:')

    p1.add_run(
                """
        1.	Electrical utilities available on site.
        2.	Sanitation facilities will be provided and available on site.
        3.	Core holes, provided by others, will be cleaned out and ready for post installation.
        4.	Fall restraint anchor points will be available and cleaned out ready for use.
        5.	Paint / PPG Duracron with a 5 year warranty.
                """)
    p1.add_run('\n')
    p1.add_run('Items EXCLUDED by Precision Rail of Oregon unless noted above:')

    p1.add_run(
        """
        1.	Deferred permits or any items not specifically included is considered furnished by others.
        2.	Taxes such as sales, local municipality, gross receipts tax and/or local business licenses.
        3.	Union, prevailing wage and/or workforce training installation
        4.	Insurance requirements above and beyond: $1M/$2M (occurrence/aggregate); and $3M
            umbrella.
        5.	Performance & payment bonds.
        6.	Pollution insurance requirements.
        7.	Deviations from project plans that impede the installation of our rail as planned.
        8.	Marking / locating rebar tensions wires
        9.	Coverage / Protection of any Glazing, Brick, Building materials
        10. Inspection for testing (example UT, NDT & others)
        11. Flaggers, and / or any personnel for traffic control
        12. Lifts, swing stages, cranes, or other equipment required to install are not included in this bid
                and are to be provided by the GC.

        """)
    p1.add_run('\n\n')

    p1.add_run("Submittal drawings with approval by the representative of buyer (customer) or owner shall be considered the correct measurement and method for fabrication. Delivery schedule will be based on receipt of final approved submittal drawings.\n\nThank you for the opportunity to submit a proposal.\n\nSincerely,")

    #rep signing
    if rep == 'jag':
        document.add_picture('JAG_signature.png', width=Inches(2))
        document.add_picture('JAG_signature.png', width=Inches(2))
        sign = document.add_paragraph()
        sign.style = document.styles['Normal']
        sign.add_run('Jeff Garlitz\n')
        sign.add_run('jgarlitz@precisionrail.com\n')
        sign.add_run('541-279-8182\n')
        
    elif rep == 'dave':
        document.add_picture('Dave_signature.png', width=Inches(2))
        document.add_picture('Dave_signature.png', width=Inches(2))
        sign = document.add_paragraph()
        sign.style = document.styles['Normal']
        sign.add_run('Dave Brown\n')
        sign.add_run('Dave@precisionrail.com\n')
        sign.add_run('503-793-1972\n')
        
    #save estimate to estimate log
    estimate_log[estimate_number] = estimate_details
    ff = open('est_log.txt','a')
    ff.write(str(estimate_log) +'\n')
    ff.close()
    
    sign.add_run('\n\n\n\n')
    sign.add_run('Acceptance of Proposal Signature _______________________              Date_______________   ')
    cwd = get_cwd()
    document.save('{}_{} - rev 0.docx'.format(job_info['company_name'],job_info['job_name']))

    #increment estimate number
    est_num = open('estimate_number.txt','w+')
    est_num.write(str(estimate_number + 1))
    est_num.close()
    
    print('proposal finished!')

    

def make_proposal():
    # create a proposal
    print('proposal writer loaded...')
    etp = Excel_to_py(file,load_workbook(filename=file,data_only=True),'Write up')
    write_proposal(etp)


def file_name():
    #open file (.xlsm)
    _path = askopenfilename(filetypes=[('Excel Files','*.xlsm')])
    print(_path)
    set_file(_path)

def get_cwd():
    # get current working directory
    cwd = os.getcwd()
    return cwd

def create_gui():
    #create tkinter gui for app
    window = Tk()
    window.title(' Auto Proposal Writer')
    window.geometry('300x100')
    open_button = Button(master=window,width=20,text='Open File',command=file_name)
    run_button =Button(master=window,width=20,text='Make Proposal',command=make_proposal)
    search_button = Button(master=window,width =15, text='search estimates',command = search_estimates)
    open_button.pack()
    run_button.pack()
    search_button.pack()
    window.mainloop()

if __name__ == '__main__':
    create_gui()
